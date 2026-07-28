#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成中文书稿排版 reference.docx（供 Pandoc --reference-doc 使用）。
规范：正文 思源宋体+Georgia 10pt 1.5倍行距 段首缩进2字符；
标题 思源宋体 黑色；A4 页边距 3.5/3.0cm；页脚居中页码。

用法:
    python reference_gen.py <project_dir>

    project_dir  书稿项目目录（如 <workspace>/projects/<project-id>）。省略时
                  从 $CLAUDE_PROJECT_DIR/projects/$BOOKFORGE_PROJECT 推导。
    产出/依赖均落在 <project_dir>/typeset/ 下（你自己的工作区），不写入插件目录：
      - 依赖 <project_dir>/typeset/base.docx
        （先跑一次 `pandoc --print-default-data-file reference.docx > base.docx` 生成）
      - 产出 <project_dir>/typeset/reference.docx
"""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def resolve_project_dir() -> str:
    if len(sys.argv) > 1:
        return sys.argv[1]
    root = os.environ.get("CLAUDE_PROJECT_DIR", os.getcwd())
    pid = os.environ.get("BOOKFORGE_PROJECT")
    if not pid:
        raise SystemExit(
            "用法: reference_gen.py <project_dir>，或设置 BOOKFORGE_PROJECT 环境变量"
        )
    return os.path.join(root, "projects", pid)


PROJECT_DIR = resolve_project_dir()
BASE = os.path.join(PROJECT_DIR, "typeset", "base.docx")
OUT = os.path.join(PROJECT_DIR, "typeset", "reference.docx")

if not os.path.exists(BASE):
    raise SystemExit(
        f"未找到 {BASE}\n"
        "先跑: pandoc --print-default-data-file reference.docx > "
        f"{BASE}"
    )

doc = Document(BASE)  # 加载 Pandoc 默认 reference（含 Body Text/First Paragraph 等全部样式）

# 页面 A4 + 页边距
sec = doc.sections[0]
sec.page_height, sec.page_width = Cm(29.7), Cm(21)
sec.top_margin = sec.bottom_margin = Cm(3.5)
sec.left_margin = sec.right_margin = Cm(3.0)


def set_font(style, latin, ea, size=None, bold=None, color=None,
             align=None, line=None, indent=None, sb=None, sa=None):
    if style is None:
        return
    rpr = style.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:ascii'), latin); rf.set(qn('w:hAnsi'), latin)
    rf.set(qn('w:eastAsia'), ea); rf.set(qn('w:cs'), latin)
    if size:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor(*color)
    pf = style.paragraph_format
    if align is not None:
        pf.alignment = align
    if line:
        pf.line_spacing = line
    if indent is not None:
        pf.first_line_indent = Pt(indent)
        # 同时设 firstLineChars（字符单位，中文规范 + officecli 识别）
        pPr = style.element.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), str(int(indent / 12 * 100)))  # 24pt→200(2字符)
    if sb is not None:
        pf.space_before = Pt(sb)
    if sa is not None:
        pf.space_after = Pt(sa)


def get_style(name):
    try:
        return doc.styles[name]
    except KeyError:
        return None


# 正文族：思源宋体 + Georgia，10pt，1.5 倍行距，段首缩进 2 字符(24pt)
BODY = dict(latin='Georgia', ea='思源宋体', size=10, line=1.5, indent=24, sa=0)
set_font(get_style('Normal'), **BODY)
set_font(get_style('Body Text'), **BODY)
set_font(get_style('First Paragraph'), **BODY)

# 标题族：思源宋体 黑色
set_font(get_style('Heading 1'), latin='Georgia', ea='思源宋体', size=22, bold=True,
         color=(0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT, indent=0, sb=24, sa=18)
set_font(get_style('Heading 2'), latin='Georgia', ea='思源宋体', size=15, bold=True,
         color=(0, 0, 0), indent=0, sb=18, sa=8)
set_font(get_style('Heading 3'), latin='Georgia', ea='思源宋体', size=12, bold=True,
         color=(0, 0, 0), indent=0, sb=12, sa=6)
set_font(get_style('Title'), latin='Georgia', ea='思源宋体', size=30, bold=True,
         color=(0, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, sb=0, sa=12)
set_font(get_style('Subtitle'), latin='Georgia', ea='思源宋体', size=15,
         color=(0, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER, indent=0)

# 页脚居中页码
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
run._r.append(f1); run._r.append(it); run._r.append(f2)

doc.save(OUT)
print(f"✓ reference.docx 已生成: {OUT}")
